from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import BinaryIO, Iterable

from docx import Document as DocxDocument
from langchain_core.documents import Document

try:
    from langchain_community.document_loaders import PyPDFLoader
except ModuleNotFoundError:
    PyPDFLoader = None


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def save_uploaded_file(
    file_obj: BinaryIO,
    file_name_or_raw_dir: str | Path,
    raw_dir: str | Path | None = None,
) -> Path:
    if raw_dir is None:
        raw_dir = file_name_or_raw_dir
        file_name = getattr(file_obj, "name", None)
        if not file_name:
            raise ValueError("Uploaded file must expose a name or file_name must be provided")
    else:
        file_name = str(file_name_or_raw_dir)

    raw_path = Path(raw_dir)
    raw_path.mkdir(parents=True, exist_ok=True)

    safe_name = _safe_file_name(file_name)
    destination = raw_path / safe_name
    _ensure_supported(destination)

    if hasattr(file_obj, "seek"):
        file_obj.seek(0)
    with destination.open("wb") as output:
        shutil.copyfileobj(file_obj, output)
    return destination


def load_document(path: str | Path) -> list[Document]:
    document_path = Path(path)
    _ensure_supported(document_path)

    if document_path.suffix.lower() == ".pdf":
        documents = _load_pdf(document_path)
    else:
        documents = [_load_docx(document_path)]

    for document in documents:
        document.metadata.update(_metadata_for(document_path))
    return documents


def load_documents(paths: Iterable[str | Path]) -> list[Document]:
    documents: list[Document] = []
    for path in paths:
        documents.extend(load_document(path))
    return documents


def _load_docx(path: Path) -> Document:
    doc = DocxDocument(path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
    return Document(page_content="\n\n".join(paragraphs), metadata={})


def _load_pdf(path: Path) -> list[Document]:
    if PyPDFLoader is not None:
        return PyPDFLoader(str(path)).load()

    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError("PDF loading requires pypdf or langchain-community.") from exc

    documents = []
    with path.open("rb") as pdf_file:
        reader = PdfReader(pdf_file)
        for page_number, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                documents.append(Document(page_content=text, metadata={"page": page_number}))
    return documents


def _safe_file_name(file_name: str) -> str:
    name = Path(file_name).name
    safe = re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("._")
    if not safe:
        raise ValueError("File name must include at least one safe character")
    return safe


def _ensure_supported(path: Path) -> None:
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type: {path.suffix or '<none>'}. Use {supported}.")


def _metadata_for(path: Path) -> dict[str, str]:
    suffix = path.suffix.lower().lstrip(".")
    return {
        "source": str(path),
        "file_name": path.name,
        "file_type": suffix,
    }
