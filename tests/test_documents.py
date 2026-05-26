from io import BytesIO

import pytest
from docx import Document as DocxDocument
from langchain_core.documents import Document

from rag_knowledge_base.documents import (
    load_document,
    load_documents,
    save_uploaded_file,
)


def test_save_uploaded_file_sanitizes_name_and_writes_content(tmp_path):
    target = save_uploaded_file(BytesIO(b"hello rag"), "../unsafe name.pdf", tmp_path)

    assert target.parent == tmp_path
    assert target.name == "unsafe_name.pdf"
    assert target.read_bytes() == b"hello rag"


def test_save_uploaded_file_rejects_unsupported_extension(tmp_path):
    with pytest.raises(ValueError, match="Unsupported file type"):
        save_uploaded_file(BytesIO(b"nope"), "notes.txt", tmp_path)


def test_load_docx_returns_langchain_documents_with_metadata(tmp_path):
    docx_path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.save(docx_path)

    documents = load_document(docx_path)

    assert len(documents) == 1
    assert isinstance(documents[0], Document)
    assert documents[0].page_content == "First paragraph\n\nSecond paragraph"
    assert documents[0].metadata["source"] == str(docx_path)
    assert documents[0].metadata["file_name"] == "sample.docx"
    assert documents[0].metadata["file_type"] == "docx"


def test_load_pdf_uses_loader_and_normalizes_metadata(tmp_path, monkeypatch):
    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n")

    class FakePdfLoader:
        def __init__(self, path):
            self.path = path

        def load(self):
            return [Document(page_content="PDF text", metadata={"page": 0})]

    monkeypatch.setattr("rag_knowledge_base.documents.PyPDFLoader", FakePdfLoader)

    documents = load_document(pdf_path)

    assert len(documents) == 1
    assert documents[0].page_content == "PDF text"
    assert documents[0].metadata["source"] == str(pdf_path)
    assert documents[0].metadata["file_name"] == "sample.pdf"
    assert documents[0].metadata["file_type"] == "pdf"
    assert documents[0].metadata["page"] == 0


def test_load_documents_flattens_supported_files(tmp_path):
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    for path, text in [(first, "One"), (second, "Two")]:
        doc = DocxDocument()
        doc.add_paragraph(text)
        doc.save(path)

    documents = load_documents([first, second])

    assert [document.page_content for document in documents] == ["One", "Two"]
